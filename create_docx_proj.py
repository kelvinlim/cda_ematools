#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Sep  7 18:29:55 2021

@author: kolim

https://python-docx.readthedocs.io/en/latest/user/quickstart.html


check all nodes causal for drinking


"""

# create a docx with a figure

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import docx
import os
import pandas as pd
import argparse
import glob
import datetime
from PIL import Image
import yaml

class PlotnWide:

    def __init__(self, mode='1wide', mdir='proj_v1_m26', 
                    esplot=True, stub = '_color.gv.png'):

    
        # need to read in config file
        self.cfgversion = 2.1
        configfile = f"config.yaml"
        self.yamlpath = os.path.join(mdir, configfile)
    
        if  not self.read_config():
            print(f"Error - config file version < {self.cfgversion}")
            exit()

        self.algo = self.cfg['causal-cmd']['algorithm']

        self.projdir = mdir
        self.proj = self.projdir.split('proj_')[1]
        self.mode = mode
        self.plotdir = f"output_{self.algo}"
        self.esplot=esplot
        self.filestub = stub

        if mode == '2wide':
            self.plot2wide()
        elif mode == '3wide':
            self.plot3wide()
        elif mode == '1wide':
            #self.plot1wide()
            self.plot()

        pass

    def create_datetime_stamp(self):
        return self.datetimestr()

    def datetimestr(self):
        dateTimeObj = datetime.datetime.now()
        timestampStr = dateTimeObj.strftime("%Y%m%d_%H%M")
        return timestampStr

    def getfiles(self, projdir='.', plotdir='output_fges',stub=''):
        # get all the paths from pathsdata.json

        if stub:
            tmpstub = stub
        else:
            tmpstub = self.filestub

        files = glob.glob(os.path.join(projdir, plotdir, f"*{tmpstub}"))
        files.sort()
        return files

    def plot(self):
        document = Document()
        fname = f'{self.proj}_plots_all_{self.create_datetime_stamp()}.docx'
        # add a header
        header = document.sections[0].header
        header.paragraphs[0].text = fname


        # keep orientation portrait
        section = document.sections[-1]
        """
        new_width, new_height = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        """
        # change margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        #cases = [584, 439, 178]
        #projdir = 'proj_' + args.proj  # set for debugging


        # get individual model plots
        files = self.getfiles(self.projdir, plotdir=self.plotdir, stub=self.filestub)

        if self.esplot:
            # add es plots to beginning
            esfiles = self.get_other_files(stub='*_effectsize_*.png')
            # place es plots files at beginning of the list
            files = esfiles + files

        total = len(files)
        count = 1

        for file in files:

            dir = self.plotdir
            
            #pngfile = os.path.join(dir, f"sub_{case:04d}_label.png")
            # use the colored lines
            pngfile = os.path.join( file)
            

            document.add_heading(f'Plot {count}/{total}  {os.path.basename(file)}')

            #document.add_picture(pngfile,width=Inches(4))
            
            document.add_picture(pngfile,
                        width=Inches(4), 
                        #height=Inches(9)
                        )



            document.add_page_break()
            count +=1

        document.save(os.path.join(self.projdir, fname))

    def get_other_files(self,stub='*_effectsize_*.png'):
        "Get the effectsize png files"
        files = glob.glob(os.path.join(self.projdir, stub))
        files.sort()
        return files

    def plot1wide(self):
        """
        Info on page orientation and margins

        https://python-docx.readthedocs.io/en/latest/user/sections.html

        plot 1 wide
        """
        document = Document()
        fname = f'{self.proj}_plots_all_{self.create_datetime_stamp()}.docx'
        # add a header
        header = document.sections[0].header
        header.paragraphs[0].text = fname

        # change orientation to landscape
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # change margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        files = self.getfiles(self.projdir)

        total = len(files)
        count = 1

        
        for file in files:

            dir = 'output_fges'

            pngfilev1 = os.path.join(file)
            pngfilev2 = pngfilev1.replace('_v1','_v2')
            pngfilev3 = pngfilev1.replace('_v1','_v3')

            pngfiles = [pngfilev1]
            table = document.add_table(rows=1, cols=len(pngfiles))

            for num in range(len(pngfiles)):

                cell = table.cell(0, num)
                paragraph = cell.paragraphs[0]
                paragraph.style = document.styles['Normal']
                
                paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify 
                run = paragraph.add_run()
                run.add_text(pngfiles[num])
                if os.path.exists(pngfiles[num]):
                    # get image dimensions
                    im = Image.open(pngfiles[num])
                    w, h = im.size
                    if h/w > 2.0:
                        run.add_picture(pngfiles[num], height=Inches(6.0))
                    else:
                        run.add_picture(pngfiles[num], width=Inches(3.0))

                #run.add_text(letters[num])
                #run.add_text("mTST="+str(tst_mean[num]))  # add tst

                #document.add_picture(pngfile,width=Inches(3.3))
                
                #document.add_picture(pngfile,width=Inches(6), height=Inches(8))


            document.add_page_break()
            count +=1

        document.save(os.path.join(self.projdir, fname))

    def plot2wide(self):
        """
        Info on page orientation and margins

        https://python-docx.readthedocs.io/en/latest/user/sections.html

        plot 2 wide
        """
        document = Document()
        fname = f'{self.proj}_plots_all_{self.create_datetime_stamp()}.docx'
        # add a header
        header = document.sections[0].header
        header.paragraphs[0].text = fname

        # change orientation to landscape
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # change margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        files = self.getfiles(self.projdir)

        total = len(files)
        count = 1

        
        for file in files:

            dir = 'output_fges'
            
            table = document.add_table(rows=1, cols=2)
            
                

            pngfilev1 = os.path.join(file)
            pngfilev2 = pngfilev1.replace('_v1','_v2')
            pngfilev3 = pngfilev1.replace('_v1','_v3')

            pngfiles = [pngfilev1, pngfilev2, pngfilev3]

            for num in range(2):

                cell = table.cell(0, num)
                paragraph = cell.paragraphs[0]
                paragraph.style = document.styles['Normal']
                
                paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify 
                run = paragraph.add_run()
                run.add_text(pngfiles[num])
                if os.path.exists(pngfiles[num]):
                    # get image dimensions
                    im = Image.open(pngfiles[num])
                    w, h = im.size
                    if h/w > 2.0:
                        run.add_picture(pngfiles[num], height=Inches(6.0))
                    else:
                        run.add_picture(pngfiles[num], width=Inches(3.0))

                #run.add_text(letters[num])
                #run.add_text("mTST="+str(tst_mean[num]))  # add tst

                #document.add_picture(pngfile,width=Inches(3.3))
                
                #document.add_picture(pngfile,width=Inches(6), height=Inches(8))


            document.add_page_break()
            count +=1

        document.save(os.path.join(self.projdir, fname))


    def plot3wide(self):
        """
        Info on page orientation and margins

        https://python-docx.readthedocs.io/en/latest/user/sections.html

        """
        document = Document()
        fname = f'{self.proj}_plots_all_{self.create_datetime_stamp()}.docx'
        # add a header
        header = document.sections[0].header
        header.paragraphs[0].text = fname

        # change orientation to landscape
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # change margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        files = self.getfiles(self.projdir)

        total = len(files)
        count = 1

        
        for file in files:

            dir = 'output_fges'
            
            table = document.add_table(rows=1, cols=3)
            
                

            pngfilev1 = os.path.join(file)
            pngfilev2 = pngfilev1.replace('_v1','_v2')
            pngfilev3 = pngfilev1.replace('_v1','_v3')

            pngfiles = [pngfilev1, pngfilev2, pngfilev3]

            for num in range(3):

                cell = table.cell(0, num)
                paragraph = cell.paragraphs[0]
                paragraph.style = document.styles['Normal']
                
                paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify 
                run = paragraph.add_run()
                run.add_text(pngfiles[num])
                if os.path.exists(pngfiles[num]):
                    run.add_picture(pngfiles[num], width=Inches(3.0))
                #run.add_text(letters[num])
                #run.add_text("mTST="+str(tst_mean[num]))  # add tst

                #document.add_picture(pngfile,width=Inches(3.3))
                
                #document.add_picture(pngfile,width=Inches(6), height=Inches(8))


            document.add_page_break()
            count +=1

        document.save(os.path.join(self.projdir, fname))

    def plot3wide_(self):
        """
        Info on page orientation and margins

        https://python-docx.readthedocs.io/en/latest/user/sections.html

        """
        document = Document()
        fname = f'{self.proj}_plots_all_{self.create_datetime_stamp()}.docx'
        # add a header
        header = document.sections[0].header
        header.paragraphs[0].text = fname

        # change orientation to landscape
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # change margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        files = self.getfiles(self.projdir)

        total = len(files)
        count = 1

        for file in files:

            dir = 'output_fges'
            
            #pngfile = os.path.join(dir, f"sub_{case:04d}_label.png")
            # use the colored lines
            pngfile = os.path.join( file)
            
            paragraph = document.add_paragraph(f'Plot {count}/{total}  {os.path.basename(file)}')

            # Creating a table object
            table = document.add_table(rows=1, cols=3)

            document.add_heading(f'Plot {count}/{total}  {os.path.basename(file)}')

            document.add_picture(pngfile,width=Inches(3.3))
            
            #document.add_picture(pngfile,width=Inches(6), height=Inches(8))


            document.add_page_break()
            count +=1

        document.save(os.path.join(self.projdir, fname))   

    def read_config(self):
        "read in the yaml config file"
        with open(self.yamlpath, 'r') as file:
            self.cfg = yaml.safe_load(file)
        
        # check version
        # does version key exist?
        if 'version' in list(self.cfg.keys()):
            if self.cfg['version'] < self.cfgversion:
                return 0
            else:
                return 1

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Plot  graphs for across  the 3 project directories.\
            Plots for same subject are placed in a single page for comparison.\
            Page is in landscape mode."
    )
    parser.add_argument("--mode", type = str,
                    help="plotting mode [1wide 2wide 3wide] def: 1wide",
                    default='1wide')
    parser.add_argument("--mdir", type = str,
                help="project directory def: proj_ior",
                default='proj_ior');
    parser.add_argument("--stub", type = str,
                help="stub of png files to include - def '_color.gv.png' ",
                default='_label.png')
    parser.add_argument('--test',  type = str,
                    help="test mode - proj_ior", 
                    default='')

    args = parser.parse_args()
    


    if args.test:

        p = PlotnWide(  mode='1wide',
                        mdir = args.test,
                        stub = args.stub,

        
        )

    else:

        p = PlotnWide(
                        mode=args.mode, 
                        mdir=args.mdir,
                        stub = args.stub
                        )